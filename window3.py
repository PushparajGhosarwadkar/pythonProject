from tkinter import *
from PIL import ImageTk, Image
import cv2
import docx
from gtts import gTTS
from playsound import playsound
from docx import Document
from docx2pdf import convert
import numpy as np
import matplotlib.pyplot as plt1
import matplotlib.pyplot as plt2
import pandas as pd
import os
import random
from playsound import playsound
import speech_recognition as sr
import threading
import tkinter
from time import sleep
import mysql.connector
import gtts

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

from docx.shared import Inches
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from datetime import datetime


# variables for important parameters

QuestionNo =2
TimeToAnswer = 4
delay = TimeToAnswer
AskHRQues = 1
Hr_Ques = 2
Cv_Ques = 2

# labels lists for answers and questions and cosine values
QuestionsList = []
ActualAnswersList = []
UserAnswersList = []
CosineList = []
AskHrQuesList = []
AskCVQuesList =[]
finalQuesList =[]

# object of docx
mydoc = docx.Document()

# importing 1st page

mydoc = docx.Document("Report.docx")
mydoc.add_page_break()
mydoc.add_heading("ClearIT : InterviewBot", 1)


# window3 settings
window3 = Tk()
window3.configure(bg='pink')
window3.geometry("{0}x{1}+0+0".format(window3.winfo_screenwidth(), window3.winfo_screenheight()))
window3.title("ClearIt Interview")

# creating oval
canvas = Canvas(window3, bg='pink')
ov = canvas.create_oval(1, 1, 120, 120, outline="red", fill="pink", width=5)
canvas.pack(fill=BOTH, expand=200)
canvas.move(ov, 1300, 580)


# Taking video
def video():
    # Create a label in the frame camera capture
    app = Frame(window3, bg="white")
    app.place(x=1100, y=30)
    lmain = Label(app)
    lmain.grid()

    # resolution 300*300
    cap = cv2.VideoCapture(0)
    cap.set(cv2.CAP_PROP_FRAME_WIDTH, 300)
    cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 300)

    def video_stream():
        _, frame = cap.read()
        cv2image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)  # conversion from brg to rgb
        img = Image.fromarray(cv2image)
        imgtk = ImageTk.PhotoImage(image=img)
        lmain.imgtk = imgtk
        lmain.configure(image=imgtk)
        lmain.after(1, video_stream)

    video_stream()


# connection to DB
mydb = mysql.connector.connect(host="localhost", user="root", passwd="12345", database="nlp")

# Question and answer words
ques = Label(window3, text="Question : ", bg='pink', font=("Times New Roman", 25)).place(x=22, y=360)
ans = Label(window3, text="Answer : ", bg='pink', font=("Times New Roman", 25)).place(x=22, y=560)

# QuestionBox
quesT = Text(window3, bg='white', relief='flat', height=5, width=120)
quesT.place(x=210, y=350)

# AnswerBox
ansT = Text(window3, bg='white', relief='flat', height=5, width=120)
ansT.place(x=210, y=550)

# interviewBot
ClearIT_text = Label(window3, text="ClearIt :", bg='pink', foreground="brown",
                     font=("Times New Roman", 50, "bold")).place(x=480, y=140)
InterviewBot_text = Label(window3, text="InterviewBot", bg='pink', foreground="brown",
                          font=("Times New Roman", 50, "bold")).place(x=580, y=220)

#adding logo
logo = Image.open("ClearIT [Logo].png")
test = ImageTk.PhotoImage(logo)
label_logo = tkinter.Label(image=test)
label_logo.image = test
label_logo.place(x=100, y=60)



# Cosine Function
def cos(Original, User):
    # Oringal answer put into variable X and User answer in variable Y
    X = str(Original)
    Y = str(User)

    # printing variable X and Variable Y
    print("Your answer is : ", X)
    print("Original answer is : ", Y)

    X_list = word_tokenize(X)  # Tokenization of X
    Y_list = word_tokenize(Y)  # Tokenization of Y

    # stopwords in english are stored in sw
    sw = stopwords.words('english')

    # 2 lists are made 1 for X and other for Y
    l1 = [];
    l2 = []

    X_set = {w for w in X_list if not w in sw}  #
    Y_set = {w for w in Y_list if not w in sw}

    rvector = X_set.union(Y_set)
    for w in rvector:
        if w in X_set:
            l1.append(1)
        else:
            l1.append(0)
        if w in Y_set:
            l2.append(1)
        else:
            l2.append(0)
    c = 0

    for i in range(len(rvector)):


        c += l1[i] * l2[i]
    cosine = c / float((sum(l1) * sum(l2)) ** 0.5)
    print("similarity: ", cosine)
    return cosine


# adding user inputs onto the report
def Userinfo():
    mycur = mydb.cursor()


    mydoc.add_paragraph(" ")
    aboutus="Our Aim : Every candidate is deserving and no one should fall back because of lack of guidance and necessary support. We here at ClearIT help students to improve their interview skills with our AI enabled InterviewBot. A very easy to use, interactive and highly capable voice bot. Our only motive is to make you professionally strong to ace in the interviews and fetch their dream job."

    mydoc.add_heading(aboutus, 2)
    mydoc.add_paragraph(" ")
    mydoc.add_heading(" ")
    quote = mydoc.add_heading("SUCCESS IS NOT FINAL;", 3)
    quote.alignment = 1
    quote2 = mydoc.add_heading("FAILURE IS NOT FATAL;", 3)
    quote2.alignment = 1
    quote3 = mydoc.add_heading("IT IS THE COURAGE TO CONTINUE", 3)
    quote3.alignment = 1
    quote4 = mydoc.add_heading("THAT COUNTS.", 3)
    quote4.alignment = 1
    mydoc.add_paragraph(" ")
    quote5 = mydoc.add_paragraph("WINSTON S. CHURCHILL", style='Quote')
    quote5.alignment = 1

    mydoc.add_heading(" ")

    name = "SELECT Name FROM tableone ORDER BY SrNo DESC LIMIT 1;"
    mycur.execute(name)
    Userinfo.data = mycur.fetchone()

    # print(data)

    name = mydoc.add_paragraph(Userinfo.data ,style = 'Title')
    name.alignment =1

    email = "SELECT email FROM tableone ORDER BY srno DESC LIMIT 1;"
    mycur.execute(email)
    data2 = mycur.fetchone()
    #print(data2)

    #email= mydoc.add_heading(data2, 3)
    email = mydoc.add_paragraph(data2, style='Title')
    email.alignment=1
    #email.style = mydoc.styles['Normal']

    department = "SELECT department FROM tableone ORDER BY srno DESC LIMIT 1;"
    mycur.execute(department)
    data3 = mycur.fetchone()
    #print(data3)

    department= mydoc.add_paragraph(data3, style='Title')
    department.alignment=1
    #department.style = mydoc.styles['Normal']

    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y, %H:%M:%S")
    date_time=str(date_time)

    timstamp = mydoc.add_paragraph(date_time, style='Title')
    timstamp.alignment = 1



    mydoc.add_page_break()
    mydoc.add_paragraph(" ")
    mydoc.add_heading(" ")
    mydoc.add_page_break()
    mydoc.add_heading("Your Responses ", 1)
    mydoc.add_paragraph(" ")
    mydoc.save("report2.docx")


Userinfo()


#Function for listening audio from user
def listen():
    #print("Say something")
    r = sr.Recognizer()
    #print("recognizing...")
    with sr.Microphone() as source:
        audio_data = r.record(source, duration=TimeToAnswer)
        text = r.recognize_google(audio_data)
        print(text)
    return text






#preparing the list of hr questions
def hrqueslist():
    for i in range(9):
        mycur = mydb.cursor()
        hrques = "SELECT HR_questions FROM hrquestions ORDER BY RAND() LIMIT 1;"
        mycur.execute(hrques)
        hrquesdata = mycur.fetchone()
        AskHrQuesList.append(hrquesdata[0])
print(AskHrQuesList)
hrqueslist()

#preparing the list of cv questions
def cvqueslist():
    mycur = mydb.cursor()
    project = "SELECT projects FROM tableone ORDER BY srno DESC LIMIT 1;"
    mycur.execute(project)
    projectdata = mycur.fetchone()
    internships = "SELECT internships FROM tableone ORDER BY srno DESC LIMIT 1;"
    mycur.execute(internships)
    internshipsdata = mycur.fetchone()
    technicalskills = "SELECT technicalskills FROM tableone ORDER BY srno DESC LIMIT 1;"
    mycur.execute(technicalskills)
    technicalskillsdata = mycur.fetchone()

    ques1 ="Tell about your role in project titled " + str(projectdata[0])
    ques2="Express your learnings from the project titled " +str(projectdata[0])
    ques3="Discuss the technologies you used for the project titled " + str(projectdata[0])
    ques4="Explain in detail about your project " +str(projectdata[0])
    ques5="Explain the real time application of " +str(projectdata[0])
    ques6="Tell me about the expierence about your internship at " + str(internshipsdata[0])
    ques7="Express your role and learnings in your internship at " +str(internshipsdata[0])
    ques8="Tell about your main accomplishments at skill " +str(technicalskillsdata[0])
    ques9="Tell something about your hobbies "
    ques10="What is the most important learning for you from engineering " + " ? "

    AskCVQuesList.extend([ques1,ques2,ques3,ques4,ques5,ques6,ques7,ques8,ques9,ques10])
    random.shuffle(AskCVQuesList)
    print(AskCVQuesList)

cvqueslist()

def questionsinreport():
    mydoc.add_page_break()
    mydoc.add_heading("Questions based on your CV", 3)
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph("Below are some more questions which can be asked to you based on your profile")
    mydoc.add_paragraph(" ")
    for i in AskCVQuesList:
        mydoc.add_paragraph(i)
    mydoc.add_page_break()
    mydoc.add_heading("Questions based on your CV", 3)
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph("Some typical HR questions which are asked in interviews")
    mydoc.add_paragraph(" ")
    for i in AskHrQuesList:
        mydoc.add_paragraph(i)
    mydoc.add_paragraph(" ")

    mydoc.save("report2.docx")

#report generation func
def report():


    #changes to be done on LEFT[],plt2.bar colours and ticklabel
    #print("this is in report function")
    #print(CosineList)
    mydoc.add_page_break()
    mydoc.add_heading("Analysis", 1)
    left = [1]
    #Height is cosinelIST

    #finding avg of cosinelist, and storing it in donut
    sum = 0
    for i in CosineList:
        sum = i + sum
    avgsum = sum * 100 / 9
    value = round(avgsum, 2)
    donutval = str(value) + "%"
    #print(donutval)
    # labels for bars
    tick_label = ['one']

    plt2.bar(left, CosineList, tick_label=tick_label, width=0.8, color=['red'])

    plt2.xlabel('Question No.')
    plt2.ylabel('Degree of correctness')

    # save the figure bar graph
    plt2.savefig('bar.png', dpi=300, bbox_inches='tight')
    plt2.close()

    #create donut fig
    name = [donutval, " "]
    size = [value, (100 - value)]
    my_circle = plt1.Circle((0, 0), 0.7, color='white')

    # Custom wedges
    plt1.pie(size, labels=name, wedgeprops={'linewidth': 7, 'edgecolor': 'white'}, colors=['green','red'])
    p = plt1.gcf()
    p.gca().add_artist(my_circle)

    # save the figure
    plt1.savefig('pie.png', dpi=300, bbox_inches='tight')
    plt1.close()
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")
    TextAnalysis= "The below Bar graphs dipicts your marks to the respective questions on a 0 to 1 scale. The x axis has the question numbers and the y value show the correcteness of your answer, where 0 is lowest and 1 is highest."

    mydoc.add_heading(TextAnalysis, 3)
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")
    mydoc.add_picture('bar.png')
    mydoc.add_page_break()

    mydoc.add_heading("Your Score", 1)
    mydoc.add_paragraph(" ")
    pieText="The below donot chart dipicts your final percentage of the interview based on your technical questions. The green area on the shows your percentage and the red is what you couldnt answer correctly"
    mydoc.add_heading(pieText,4)
    im = mydoc.add_paragraph()
    im.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = im.add_run()
    mydoc.add_paragraph(" ")

    run.add_picture('pie.png')
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph("Your final score : "+donutval,style='TOC Heading')
    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")

    if (0 < avgsum < 35):
        bad = "Your performance in the mock interview is below average. You need to put in serious efforts to pull yourself up. Have good grasp on fundamental concepts and make your base strong."
        mydoc.add_heading(bad, 3)
        mydoc.add_paragraph(" ")
        para1 = mydoc.add_heading("ALL THE VERY BEST!!!", 3)
        para1.alignment = 1

    elif (35 < avgsum < 75):
        average = "Your performance in the mock interview was satisfactory. You have the capacity to take yourself ahead and become above average candidate. Be confident while answering. Keep the pace on and you will clear the interviews."
        mydoc.add_heading(average, 3)
        mydoc.add_paragraph(" ")
        para2 = mydoc.add_heading("ALL THE VERY BEST!!!", 3)
        para2.alignment = 1
    else:
        excellent = "Your performance in the mock interview was impressive. You are a really good candidate to be onboarded. Have confidence in yourself and you will fetch your dream job."
        mydoc.add_heading(excellent, 3)
        mydoc.add_paragraph(" ")
        para3 = mydoc.add_heading("ALL THE VERY BEST !!", 3)
        para3.alignment = 1



    mydoc.add_page_break()
    questionsinreport()
    mydoc.add_page_break()
    mydoc.add_heading("DO'S & DON'TS In an Interview", 1)
    mydoc.add_paragraph(" ")
    mydoc.add_picture("image.png", width=Inches(7), height=Inches(8))
    mydoc.add_page_break()
    mydoc.add_heading("Thank You", 1)
    mydoc.add_paragraph(" ")

    thanks = "Thank you for opting ClearIT Mock interviews. We hope our automated InterviewBot and interview analysis report polished your skills and helped you to identify loop holes. We are sure you will reach your full potency. Keep the pace on and you will definitely succeed. "
    mydoc.add_heading(thanks, 3)
    mydoc.add_paragraph(" ")
    abc = mydoc.add_heading("Give your best forget the rest !!!", 3)
    abc.alignment = 1
    mydoc.add_paragraph(" ")
    mydoc.add_heading(
        "If our InterviewBot helped you, we urge you to recommend it to others too. Spread the possibility.", 3)

    mydoc.add_paragraph(" ")
    mydoc.add_paragraph(" ")
    luck = mydoc.add_paragraph()
    luck.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = luck.add_run()
    run.add_picture('luck.jpeg', width=Inches(4), height=Inches(3))

    footer_section = mydoc.sections[0]
    footer = footer_section.footer
    footer_text = footer.paragraphs[0]
    footer_text.text = "For any queries or concerns contact: beprojectece@gmail.com"
    mydoc.save("report2.docx")

# CountDown Function'
counter = Label(window3, bg='pink', fg='red', font=("Times New Roman", 60, "bold", "italic"))
counter.place(x=1335, y=590)
def countdown(count):
    counter['text'] = count
    if count > 0:
        window3.after(1000, countdown, count - 1)

i = 0
k=0
# playing text as audio
def mainloop():

    intro = "Hello" + " " + Userinfo.data[0] + "," + " " + "Please Introduce yourself"
    IntroLabel = Label(window3, text=str(intro), bg='white', font=("Times New Roman", 25))  # prints question
    IntroLabel.place(x=210, y=350)
    tts = gtts.gTTS(str(intro))  # test to speech progress
    tts.save("intro.mp3")  # saves tts file
    playsound("intro.mp3")  # plays the tts
    #print(str(labels[i]))
    if("intro.mp3"):
        countdown(TimeToAnswer)
    UserAnswer = listen()
    AnsLabel = Label(window3, text=UserAnswer, bg='white', font=("Times New Roman", 25))
    AnsLabel.place(x=220, y=560)

    os.remove("intro.mp3")
    window3.after(TimeToAnswer * 1000, lambda: IntroLabel.destroy())
    window3.after(TimeToAnswer * 1000, lambda: AnsLabel.destroy())
    sleep(5)


    for x in range(QuestionNo):
        mycursor = mydb.cursor()
        query = "SET @result= (SELECT Questions FROM questionsdb ORDER BY RAND() LIMIT 1);"
        mycursor.execute(query)
        query2 = "SELECT @result;"
        mycursor.execute(query2)
        ques = mycursor.fetchone()
        QuestionsList.append(ques[0])
        print("Question : ",ques[0])
        SelectAnswerQuery = "SELECT ActualAnswer FROM questionsdb WHERE Questions= @result;"
        mycursor.execute(SelectAnswerQuery)
        ActualAns = mycursor.fetchone()
        print( "correct answer is : ", ActualAns[0])
        #ActualAns = str(ans)
        ActualAnswersList.append(ActualAns[0])
        #print("Actual answer", ActualAns[0])
        #print(ques)
    for y in range(Cv_Ques):
        QuestionsList.append(AskCVQuesList[y])
    for z in range(Hr_Ques):
        QuestionsList.append(AskHrQuesList[z])

    QuestionsList.append("Interview is over please press next")

    def update():

        global i
        if (i <= len(QuestionsList)):
            global CosineList
            QuesLabel = Label(window3, text=QuestionsList[i], bg='white', font=("Times New Roman", 25))  # prints question
            QuesLabel.place(x=210, y=350)
            tts = gtts.gTTS(str(QuestionsList[i]))  # test to speech progress
            tts.save("ques.mp3")  # saves tts file
            playsound("ques.mp3")  # plays the tts
            # print(str(labels[i]))
            if(i<=len(QuestionsList)):
                if ("ques.mp3" and QuesLabel):
                    countdown(TimeToAnswer)
                os.remove("ques.mp3")
                UserAnswer = listen()
                AnsLabel = Label(window3, text=UserAnswer, bg='white', font=("Times New Roman", 25))
                AnsLabel.place(x=220, y=560)
                global k
                if(k<QuestionNo):
                    UserAnswer = str(UserAnswer)
                    UserAnswersList.append(UserAnswer)
                    CosineVal=cos(UserAnswer, ActualAns) #cos value stored in CosineVal
                    CosineString=str(CosineVal)
                    CosineList.append(CosineVal+0.15)
                    mydoc.add_heading("Question", 3)
                    mydoc.add_paragraph(QuestionsList[i])  # adding questions to word document
                    mydoc.add_heading("Expected Answer : ", 3)
                    mydoc.add_paragraph(ActualAnswersList[k])
                    mydoc.add_heading("Your Answer : ", 3)
                    mydoc.add_paragraph(UserAnswersList[k])
                    mydoc.save("report2.docx")
                    k=k+1
                i = i + 1
                window3.after((TimeToAnswer) * 1000, lambda: QuesLabel.destroy())
                window3.after((TimeToAnswer) * 1000, lambda: AnsLabel.destroy())
                if(i<=len(QuestionsList) and TimeToAnswer+4):
                    update()

    update() # recalls the update function

# multitreading buttons
#video = tkinter.Button(window3, text="", background="pink", relief="flat", height=1, width=1, fg="white", command=threading.Thread(target=video).start()).place(x=400, y=500)

Mainloop = tkinter.Button(window3, text="", background="pink", relief="flat", height=1, width=1, fg="white",command=threading.Thread(target=mainloop).start()).place(x=400, y=500)

def window4():
    report()
    os.system('window4.py')

# next button
next4 = Button(window3, height=2, width=7, text="Next", background="orange", fg="white",command=lambda: threading.Thread(target=window4).start())
next4.place(x=950, y=700)
next4.config(state='normal')


#generatereport = Button(window3, height=2, width=14, text="Generate Report", background="orange", fg="white",command=lambda: threading.Thread(target=report()).start())
#generatereport.place(x=750, y=700)
#generatereport.config(state='disable')

# quit button
quit = Button(window3, text="Quit", background="Red", height=2, width=7, fg="white", command=window3.destroy)
quit.place(x=500, y=700)

window3.mainloop()